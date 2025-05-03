import streamlit as st
from openai import OpenAI
import os
import tempfile
import time
from pathlib import Path
import json
import re
from concurrent.futures import ThreadPoolExecutor

# Bibliothèques pour traiter différents types de documents
import PyPDF2
import docx
import io

# Configuration de la page Streamlit
st.set_page_config(
    page_title="Assistant IA - Dialogue & Q&A sur Documents",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Style CSS personnalisé pour améliorer l'interface
st.markdown("""
<style>
.chat-message {
    padding: 1.5rem;
    border-radius: 0.5rem;
    margin-bottom: 1rem;
    display: flex;
    flex-direction: row;
    align-items: flex-start;
}
.chat-message.user {
    background-color: #f0f2f6;
}
.chat-message.assistant {
    background-color: #e3f2fd;
}
.chat-message .avatar {
    width: 40px;
    height: 40px;
    border-radius: 50%;
    object-fit: cover;
    margin-right: 1rem;
}
.chat-message .message {
    flex-grow: 1;
}
.chat-message .doc-indicator {
    font-size: 0.8rem;
    color: #424242;
    margin-top: 0.5rem;
}
.document-pill {
    display: inline-block;
    background-color: #e0e0e0;
    padding: 0.2rem 0.5rem;
    border-radius: 1rem;
    font-size: 0.8rem;
    margin-right: 0.5rem;
    margin-top: 0.3rem;
}
.stTextInput div[data-baseweb="base-input"] {
    display: flex;
    flex-direction: row;
    align-items: center;
}
.upload-button {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    cursor: pointer;
    padding: 0.25rem;
    border-radius: 0.25rem;
    margin-right: 0.5rem;
}
.upload-button:hover {
    background-color: #f0f2f6;
}
/* Styles pour les avatars améliorés - version plus générique */
.avatar-user {
    background-color: #7b95b4;
    color: white;
    font-weight: bold;
    display: flex;
    align-items: center;
    justify-content: center;
    border-radius: 50%;
    width: 40px;
    height: 40px;
}
.avatar-assistant {
    background-color: #2ecc71;
    color: white;
    font-weight: bold;
    display: flex;
    align-items: center;
    justify-content: center;
    border-radius: 50%;
    width: 40px;
    height: 40px;
}
</style>
""", unsafe_allow_html=True)

# Fonction de cache pour éviter de recalculer des résultats déjà obtenus
@st.cache_data(ttl=3600)
def get_api_credentials():
    """Récupère les identifiants d'API avec mise en cache"""
    return {
        "base_url": os.environ.get("SCALEWAY_API_BASE_URL"),
        "api_key": os.environ.get("SCALEWAY_API_KEY")
    }

# Configuration de l'API Scaleway
api_creds = get_api_credentials()
API_BASE_URL = api_creds["base_url"]
API_KEY = api_creds["api_key"]

if not API_BASE_URL or not API_KEY:
    st.error("Les variables d'environnement SCALEWAY_API_BASE_URL et SCALEWAY_API_KEY doivent être définies.")
    st.stop()

# Paramètres du modèle
MODEL = "llama-3.3-70b-instruct"
MAX_TOKENS = 4096
TEMPERATURE = 0.6
TOP_P = 0.9
PRESENCE_PENALTY = 0.0
STOP_SEQUENCE = ["/stop"]

# Initialisation des variables de session avec un mécanisme plus robuste
def init_session_state():
    """Initialise les variables de session de façon plus structurée"""
    if 'initialized' not in st.session_state:
        st.session_state.conversation_history = [
            {"role": "system", "content": "Tu es un assistant intelligent qui répond en français même si la question est dans une autre langue. Tu peux discuter de tout sujet et analyser des documents si l'utilisateur en fournit."}
        ]
        st.session_state.chat_messages = []
        st.session_state.documents = {}  # Dictionnaire pour stocker {nom_document: contenu}
        st.session_state.submitted = False
        st.session_state.initialized = True

# Appel de l'initialisation
init_session_state()

# Fonctions pour extraire le texte de différents formats de documents
def extract_text_from_pdf(file):
    """Extrait le texte d'un fichier PDF avec une gestion d'erreurs améliorée et robuste"""
    text = ""
    temp_file_path = None
    
    try:
        # Créer un fichier temporaire avec un meilleur contrôle
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(file.getvalue())
            temp_file_path = temp_file.name
        
        # Ouvrir et lire le PDF avec gestion des erreurs améliorée
        try:
            with open(temp_file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Vérifier si le PDF est crypté et essayer de le décrypter si nécessaire
                if pdf_reader.is_encrypted:
                    try:
                        # Essayer avec un mot de passe vide (beaucoup de PDFs sont marqués comme cryptés mais sans mot de passe)
                        pdf_reader.decrypt('')
                    except:
                        st.warning("Le PDF semble être protégé par un mot de passe et ne peut pas être entièrement analysé.")
                
                # Extraction page par page avec gestion des erreurs pour chaque page
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                        else:
                            # Essayer une méthode alternative pour les PDFs problématiques
                            st.info(f"Méthode alternative d'extraction utilisée pour la page {page_num+1}")
                    except Exception as page_error:
                        st.warning(f"Impossible d'extraire le texte de la page {page_num+1}: {str(page_error)}")
                        # Continuer avec les autres pages
        
        except PyPDF2.errors.PdfReadError as pdf_error:
            st.error(f"Erreur lors de la lecture du PDF: {str(pdf_error)}")
            st.info("Tentative avec une méthode alternative...")
            
            # Méthode alternative en cas d'échec de PyPDF2
            try:
                import pdfplumber
                
                with pdfplumber.open(temp_file_path) as pdf:
                    for page in pdf.pages:
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n\n"
                        except:
                            pass  # Ignorer les pages problématiques
            except ImportError:
                st.error("Module pdfplumber non disponible pour l'extraction alternative.")
                # Recommander l'installation: pip install pdfplumber
    
    except Exception as e:
        st.error(f"Erreur lors de l'extraction du texte du PDF: {str(e)}")
    
    finally:
        # Nettoyer le fichier temporaire
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception as cleanup_error:
                pass  # Ignorer les erreurs de nettoyage
    
    # Vérifier si du texte a été extrait
    if not text.strip():
        st.warning("Aucun texte n'a pu être extrait du PDF. Cela peut être dû à un PDF scanné ou protégé.")
    
    return text

def extract_text_from_docx(file):
    """Extrait le texte d'un fichier DOCX avec gestion d'erreurs améliorée"""
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file.getvalue()))
        # Extraction du texte des paragraphes et des tableaux
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extraction du texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
    except Exception as e:
        st.error(f"Erreur lors de l'extraction du texte du DOCX: {str(e)}")
    
    return text

def extract_text_from_txt(file_object):
    """Extrait le texte d'un fichier TXT avec gestion d'erreurs améliorée"""
    try:
        # Si c'est un objet UploadedFile (de l'interface Streamlit)
        if hasattr(file_object, 'getvalue'):
            return file_object.getvalue().decode("utf-8")
        # Si c'est un objet BufferedReader (fichier ouvert)
        else:
            file_object.seek(0)  # Retour au début du fichier
            content = file_object.read()
            if isinstance(content, bytes):
                return content.decode("utf-8")
            return content
    except UnicodeDecodeError:
        # Essaie avec différents encodages si UTF-8 échoue
        encodings = ['latin-1', 'iso-8859-1', 'windows-1252']
        for encoding in encodings:
            try:
                if hasattr(file_object, 'getvalue'):
                    return file_object.getvalue().decode(encoding)
                else:
                    file_object.seek(0)
                    content = file_object.read()
                    if isinstance(content, bytes):
                        return content.decode(encoding)
                    return content
            except UnicodeDecodeError:
                continue
        st.error("Impossible de déterminer l'encodage du fichier texte.")
        return ""
    except Exception as e:
        st.error(f"Erreur lors de l'extraction du texte du fichier TXT: {str(e)}")
        return ""

@st.cache_data(ttl=3600, show_spinner=False)
def process_file(file_content, file_name):
    """Traite le fichier uploadé et extrait son contenu textuel (version avec cache)"""
    file_extension = Path(file_name).suffix.lower()
    
    # Création d'un fichier temporaire avec le contenu
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        temp_file.write(file_content)
        temp_path = temp_file.name
    
    try:
        if file_extension == '.pdf':
            with open(temp_path, 'rb') as f:
                uploaded_file = type('', (), {})()  # Crée un objet vide
                uploaded_file.getvalue = lambda: file_content  # Ajoute une méthode getvalue
                result = extract_text_from_pdf(uploaded_file)
        elif file_extension == '.docx':
            with open(temp_path, 'rb') as f:
                uploaded_file = type('', (), {})()  # Crée un objet vide
                uploaded_file.getvalue = lambda: file_content  # Ajoute une méthode getvalue
                result = extract_text_from_docx(uploaded_file)
        elif file_extension == '.txt':
            with open(temp_path, 'rb') as f:
                result = extract_text_from_txt(f)
        else:
            st.error(f"Format de fichier non pris en charge: {file_extension}")
            result = ""
    finally:
        # Nettoyage du fichier temporaire
        try:
            os.unlink(temp_path)
        except:
            pass
            
    return result

def get_chunks(text, chunk_size=3000, overlap=200):
    """Divise le texte en chunks pour gérer les documents longs"""
    chunks = []
    if not text:
        return chunks
        
    # Chunking amélioré avec détection de phrases
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text) and end - start == chunk_size:
            # Chercher la fin de phrase/paragraphe la plus proche pour une coupure propre
            # Recherche plus sophistiquée avec pattern de fin de phrase
            sentence_end_pattern = re.compile(r'[.!?]\s+')
            matches = list(sentence_end_pattern.finditer(text[end-200:end]))
            if matches:
                # Utilise la dernière correspondance trouvée
                last_match = matches[-1]
                end = end - 200 + last_match.end()
            else:
                # Cherche une fin de ligne si pas de fin de phrase
                newline_matches = list(re.finditer(r'\n', text[end-200:end]))
                if newline_matches:
                    last_newline = newline_matches[-1]
                    end = end - 200 + last_newline.end()
                    
        chunks.append(text[start:end])
        start = end - overlap if end < len(text) else end
    
    return chunks

def create_context_for_question(question, documents, max_length=6000):
    """Crée un contexte pertinent pour la question en utilisant les documents disponibles"""
    if not documents:
        return ""
    
    # Combine tous les documents en un seul texte pour l'analyse
    all_text = ""
    for doc_name, doc_content in documents.items():
        all_text += f"\n\n--- DOCUMENT: {doc_name} ---\n\n"
        all_text += doc_content
    
    # Si le texte total est petit, on utilise tout
    if len(all_text) <= max_length:
        return all_text
    
    # Pour les documents plus grands, on utilise une méthode de recherche plus sophistiquée
    chunks = get_chunks(all_text)
    
    # Analyse sémantique améliorée de la question
    # Extraction des mots-clés avec élimination des stopwords
    words = re.findall(r'\b\w+\b', question.lower())
    stopwords = set(['le', 'la', 'les', 'un', 'une', 'des', 'et', 'est', 'à', 'au', 'aux', 
                    'de', 'du', 'en', 'ce', 'cette', 'ces', 'qui', 'que', 'quoi', 'où', 
                    'comment', 'pourquoi', 'quand', 'quel', 'quelle', 'quels', 'quelles',
                    'il', 'elle', 'ils', 'elles', 'nous', 'vous', 'leur', 'leurs', 'son',
                    'sa', 'ses', 'mon', 'ma', 'mes', 'ton', 'ta', 'tes', 'pour', 'par',
                    'avec', 'sans', 'mais', 'ou', 'où', 'donc', 'or', 'ni', 'car', 'sur'])
    
    keywords = [word for word in words if word not in stopwords and len(word) > 2]
    
    if not keywords:
        # Si pas de mots-clés significatifs, on prend les premiers chunks
        selected_chunks = list(range(min(5, len(chunks))))
    else:
        # Score chaque chunk en fonction de la pertinence
        chunk_scores = []
        for i, chunk in enumerate(chunks):
            chunk_lower = chunk.lower()
            
            # Score basé sur la fréquence des mots-clés avec pondération
            base_score = 0
            for keyword in keywords:
                # Donne un poids plus élevé aux mots plus longs (supposés plus significatifs)
                weight = min(1.0, 0.5 + (len(keyword) / 10))
                count = chunk_lower.count(keyword)
                base_score += count * weight
            
            # Bonus pour les chunks contenant des phrases complètes de la question
            question_phrases = re.split(r'[.!?]', question.lower())
            phrase_bonus = 0
            for phrase in question_phrases:
                if len(phrase.strip()) > 10 and phrase.strip() in chunk_lower:
                    phrase_bonus += 2
            
            final_score = base_score + phrase_bonus
            chunk_scores.append((i, final_score))
        
        # Trie les chunks par score et prend les meilleurs jusqu'à atteindre max_length
        sorted_chunks = sorted(chunk_scores, key=lambda x: x[1], reverse=True)
        
        # Sélectionne les chunks avec le meilleur score
        selected_chunks = []
        total_length = 0
        
        for chunk_idx, score in sorted_chunks:
            if score > 0 and total_length + len(chunks[chunk_idx]) <= max_length:
                selected_chunks.append(chunk_idx)
                total_length += len(chunks[chunk_idx])
        
        # Si aucun chunk n'a de score positif ou si on n'a pas assez de contenu
        if not selected_chunks or total_length < max_length * 0.5:
            # Ajoute des chunks supplémentaires au début du document
            for i in range(min(3, len(chunks))):
                if i not in selected_chunks and total_length + len(chunks[i]) <= max_length:
                    selected_chunks.append(i)
                    total_length += len(chunks[i])
    
    # Trie les indices pour préserver l'ordre original des documents
    selected_chunks.sort()
    context = "\n\n".join([chunks[i] for i in selected_chunks])
    
    return context

def add_message(role, content, attached_docs=None):
    """Ajoute un message à l'interface de chat et à l'historique de conversation"""
    # Ajoute à l'état de session pour l'affichage
    st.session_state.chat_messages.append({
        "role": role, 
        "content": content,
        "attached_docs": attached_docs
    })
    
    # Ajoute à l'historique de conversation pour le LLM
    st.session_state.conversation_history.append({
        "role": role,
        "content": content
    })

# Fonction pour afficher les messages de chat avec un style amélioré
def display_messages():
    for idx, msg in enumerate(st.session_state.chat_messages):
        role = msg["role"]
        content = msg["content"]
        attached_docs = msg.get("attached_docs", None)
        
        if role == "user":
            # Utilisation d'une icône plus générique pour l'utilisateur
            avatar_html = '<div class="avatar-user">👤</div>'
            bg_color = "user"
        else:
            # Icône pour l'assistant
            avatar_html = '<div class="avatar-assistant">🤖</div>'
            bg_color = "assistant"
        
        # Construit l'affichage du message avec markdown pour le contenu
        message_html = f"""
        <div class="chat-message {bg_color}" id="message-{idx}">
            {avatar_html}
            <div class="message">
                {content}
        """
        
        # Ajoute des indicateurs pour les documents attachés
        if attached_docs:
            message_html += '<div class="doc-indicator">Documents attachés: '
            for doc in attached_docs:
                message_html += f'<span class="document-pill">📄 {doc}</span>'
            message_html += '</div>'
        
        message_html += "</div></div>"
        
        st.markdown(message_html, unsafe_allow_html=True)

# Cache du client OpenAI pour éviter de le recréer à chaque interaction
@st.cache_resource
def get_openai_client():
    """Récupère un client OpenAI avec mise en cache"""
    return OpenAI(
        base_url=API_BASE_URL,
        api_key=API_KEY
    )

# Interface utilisateur Streamlit optimisée
# Remplacez la partie initiale du code main() par cette version

def main():
    # Sidebar pour les paramètres
    with st.sidebar:
        st.title("⚙️ Paramètres")
        
        # Affichage des documents déjà téléchargés
        if st.session_state.documents:
            st.write("### Documents disponibles:")
            for doc_name in st.session_state.documents.keys():
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"📄 {doc_name}")
                with col2:
                    if st.button("❌", key=f"delete_{doc_name}"):
                        del st.session_state.documents[doc_name]
                        st.success(f"Document '{doc_name}' supprimé")
                        st.rerun()
        
        # Paramètres avancés
        with st.expander("⚙️ Paramètres avancés", expanded=False):
            st.slider("Température", min_value=0.0, max_value=1.0, value=TEMPERATURE, step=0.1, key="temperature", 
                      help="Contrôle la créativité des réponses (0=déterministe, 1=créatif)")
            st.slider("Longueur maximale", min_value=100, max_value=4096, value=MAX_TOKENS, step=100, key="max_tokens",
                      help="Nombre maximum de tokens dans la réponse")
            
            # Option pour télécharger l'historique de conversation
            if st.button("Télécharger l'historique"):
                conversation_json = json.dumps(st.session_state.chat_messages, ensure_ascii=False, indent=2)
                st.download_button(
                    label="Télécharger JSON",
                    data=conversation_json,
                    file_name="conversation_history.json",
                    mime="application/json",
                )
            
            # Bouton de réinitialisation dans les paramètres avancés
            if st.button("Réinitialiser la conversation", key="reset_conversation_advanced"):
                st.session_state.conversation_history = [
                    {"role": "system", "content": "Tu es un assistant intelligent qui répond en français même si la question est dans une autre langue. Tu peux discuter de tout sujet et analyser des documents si l'utilisateur en fournit."}
                ]
                st.session_state.chat_messages = []
                st.session_state.documents = {}
                st.success("Conversation réinitialisée!")
                st.rerun()

    # En-tête avec titre et bouton Nouvelle Conversation
    col1, col2 = st.columns([5, 1])
    with col1:
        st.title("🧠 Assistant IA - Dialogue & Documents")
    with col2:
        # Bouton Nouvelle Conversation simple et direct avec Streamlit
        if st.button("➕ Nouvelle", key="new_conversation", use_container_width=True):
            st.session_state.conversation_history = [
                {"role": "system", "content": "Tu es un assistant intelligent qui répond en français même si la question est dans une autre langue. Tu peux discuter de tout sujet et analyser des documents si l'utilisateur en fournit."}
            ]
            st.session_state.chat_messages = []
            st.session_state.documents = {}
            st.success("Nouvelle conversation démarrée!")
            st.rerun()
    
    # Style CSS pour le bouton Nouvelle
    st.markdown("""
    <style>
        /* Style pour le bouton Nouvelle Conversation */
        button[data-testid="baseButton-secondary"]:has(div:contains("➕ Nouvelle")) {
            background-color: #4CAF50 !important;
            color: white !important;
            font-weight: bold !important;
            border-radius: 20px !important;
            border: none !important;
            box-shadow: 0 2px 5px rgba(0,0,0,0.2) !important;
            transition: all 0.3s ease !important;
        }
        
        button[data-testid="baseButton-secondary"]:has(div:contains("➕ Nouvelle")):hover {
            background-color: #45a049 !important;
            box-shadow: 0 4px 8px rgba(0,0,0,0.3) !important;
            transform: translateY(-2px) !important;
        }
        
        /* Style pour les notifications de succès */
        div[data-testid="stSuccessMessage"] {
            position: fixed !important;
            top: 20px !important;
            right: 20px !important;
            z-index: 9999 !important;
            animation: fadeOut 3s forwards !important;
            animation-delay: 2s !important;
        }
        
        @keyframes fadeOut {
            from { opacity: 1; }
            to { opacity: 0; visibility: hidden; }
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.write("Discutez avec l'assistant et attachez des documents au besoin pour poser des questions dessus.")
    
    # Affichage des messages de chat
    with st.container():
        display_messages()
    
    # Zone de saisie pour la question
    st.write("### Envoyez un message")
    
    # Layout pour le champ de saisie et l'upload de fichier
    col1, col2 = st.columns([6, 1])
    
    # Gestion du fichier uploader
    with col2:
        uploaded_files = st.file_uploader(
            "Joindre un document", 
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            label_visibility="collapsed",
            key="file_upload"
        )
    
    # Indication du raccourci clavier
    st.markdown("""
    <div style="text-align: right; font-size: 0.8em; color: #888; margin-top: -15px; margin-bottom: 5px;">
        Envoyez avec Cmd+Enter (Mac) ou Ctrl+Enter (PC)
    </div>
    """, unsafe_allow_html=True)
    
    # Gestion des réinitialisations du champ de saisie
    if 'clear_input' in st.session_state and st.session_state.clear_input:
        st.session_state.user_input = ""
        st.session_state.clear_input = False
    
    # Champ de saisie utilisateur
    with col1:
        user_input = st.text_area("Votre message:", height=100, key="user_input", 
                                 label_visibility="collapsed")
    
    # Traitement des fichiers téléchargés
    attached_docs = []
    
    if uploaded_files:
        st.write("Documents à joindre à ce message:")
        cols = st.columns(4)
        
        for i, uploaded_file in enumerate(uploaded_files):
            col_index = i % 4
            with cols[col_index]:
                file_name = uploaded_file.name
                st.write(f"📎 {file_name}")
                
                # Traitement du fichier
                with st.spinner(f"Traitement de {file_name}..."):
                    document_text = process_file(uploaded_file.getvalue(), file_name)
                    if document_text:
                        st.session_state.documents[file_name] = document_text
                        attached_docs.append(file_name)
                        st.success(f"✓ ({len(document_text)} caractères)")
                    else:
                        st.error("Échec du traitement")
    
    # Bouton d'envoi
    send_button = st.button("Envoyer", key="send_message")
    
    # Détection de l'envoi (bouton ou Cmd+Enter)
    send_pressed = send_button or (user_input and user_input.endswith('\n'))
    
    if send_pressed:
        # Vérification qu'il y a un message ou un document
        if not user_input.strip() and not attached_docs:
            st.warning("Veuillez entrer un message ou joindre un document.")
        else:
            # Nettoie l'entrée utilisateur
            cleaned_input = user_input.strip()
            
            # Cas spécial: document sans message
            if not cleaned_input and attached_docs:
                cleaned_input = "Peux-tu analyser ce(s) document(s) et me dire ce qu'il(s) contien(nen)t?"
            
            # Ajoute le message utilisateur à l'interface
            add_message("user", cleaned_input, attached_docs)
            
            # Préparation du contexte des documents
            document_context = ""
            if attached_docs:
                # Utilise seulement les documents joints à ce message
                docs_for_context = {name: content for name, content in st.session_state.documents.items() 
                                    if name in attached_docs}
                with st.spinner("Analyse des documents..."):
                    document_context = create_context_for_question(cleaned_input, docs_for_context)
            
            # Préparation des messages pour l'API
            if document_context:
                # Message système pour mode documents
                system_message = {"role": "system", "content": "Tu es un assistant intelligent qui répond en français. Tu peux analyser des documents fournis par l'utilisateur et répondre à des questions à leur sujet."}
                
                # Liste des messages pour l'API (tous sauf le dernier)
                messages = [system_message]
                for msg in st.session_state.conversation_history[1:-1]:  # Exclut system et dernier message
                    messages.append(msg)
                
                # Création du prompt avec contexte document
                full_prompt = f"""Voici ma question: {cleaned_input}

Je joins également les documents suivants pour référence:

{document_context}

Réponds à ma question en te basant sur les informations fournies dans ces documents si pertinent."""
                
                # Ajoute le prompt enrichi à la liste des messages
                messages.append({"role": "user", "content": full_prompt})
            else:
                # Mode conversation normal sans documents
                system_message = {"role": "system", "content": "Tu es un assistant intelligent qui répond en français même si la question est dans une autre langue."}
                messages = [system_message] + st.session_state.conversation_history[1:]
            
            # Récupération du client OpenAI
            client = get_openai_client()
            
            # Génération de la réponse
            with st.spinner("Génération de la réponse..."):
                full_response = ""
                
                try:
                    # Appel API en streaming
                    response = client.chat.completions.create(
                        model=MODEL,
                        messages=messages,
                        max_tokens=st.session_state.get("max_tokens", MAX_TOKENS),
                        temperature=st.session_state.get("temperature", TEMPERATURE),
                        top_p=TOP_P,
                        presence_penalty=PRESENCE_PENALTY,
                        stop=STOP_SEQUENCE,
                        stream=True,
                    )
                    
                    # Affichage en streaming
                    with st.container():
                        message_container = st.empty()
                        
                        for chunk in response:
                            if chunk.choices and chunk.choices[0].delta.content:
                                content = chunk.choices[0].delta.content
                                full_response += content
                                
                                # Message avec icône générique
                                message_html = f"""
                                <div class="chat-message assistant">
                                    <div class="avatar-assistant">🤖</div>
                                    <div class="message">
                                        {full_response}
                                    </div>
                                </div>
                                """
                                message_container.markdown(message_html, unsafe_allow_html=True)
                                time.sleep(0.01)
                    
                    # Ajout de la réponse à l'historique
                    add_message("assistant", full_response)
                    
                    # Indique que le champ de saisie doit être vidé
                    st.session_state.clear_input = True
                    
                    # Supprime les données du file uploader
                    if "file_upload" in st.session_state:
                        del st.session_state.file_upload
                    
                    # Rafraîchit l'interface
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Erreur lors de la génération de la réponse: {str(e)}")
                    st.error(f"Détails de l'erreur: {type(e).__name__}")

# Point d'entrée de l'application
if __name__ == "__main__":
    main()
